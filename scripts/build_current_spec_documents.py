from __future__ import annotations

import shutil
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Flowable,
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "output" / "docx"
OUT_PDF = ROOT / "output" / "pdf"
OUT_RENDER = ROOT / "output" / "rendered-spec"
DOCX_PATH = OUT_DOCX / "selenia-aroma-current-spec-with-screenshots-2026-07-30.docx"
PDF_PATH = OUT_PDF / "selenia-aroma-current-spec-with-screenshots-2026-07-30.pdf"
SCREENSHOT_DIR = ROOT / "output" / "screenshots" / "customer-public"

SCREENSHOTS = [
    ("ホーム", "/dashboard", "01-dashboard.png", "行動ハブとして、あいさつ、おすすめカード、次アクションを表示。"),
    ("アロマ一覧", "/aromas", "02-aromas.png", "購入者自身のアロマ記録をカード形式で一覧表示。"),
    ("アロマ詳細", "/aromas/relax-night", "03-aroma-detail.png", "テーマ、ベースブレンド、追加オイル名、注意事項、再購入導線を表示。"),
    ("再購入確認", "/aromas/relax-night/reorder", "04-reorder.png", "外部ECサイトへ移動する前の確認画面。"),
    ("気分から探す", "/moods", "05-moods.png", "気分・目的別の入口をカードで表示。"),
    ("気分別詳細", "/moods/relax", "06-mood-relax.png", "選択した目的に合う一般精油を表示。"),
    ("ベースブレンド一覧", "/base-blends", "07-base-blends.png", "12種類のベースブレンドを図鑑形式で表示。"),
    ("ベースブレンド詳細", "/base-blends/base-01", "08-base-blend-detail.png", "配合比率を伏せ、含まれる精油名と目的を表示。"),
    ("精油詳細", "/oils/lavender", "09-oil-lavender.png", "一般精油の説明、香りの印象、使用シーン、相性を表示。"),
    ("プロフィール", "/profile", "10-profile.png", "会員情報、お気に入り、好みの香りタイプを表示。"),
]

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 35, 35)
MUTED = RGBColor(90, 90, 90)
TABLE_FILL = "F2F4F7"
CALLOUT_FILL = "F6F2FB"
CALLOUT_BORDER = "D8C8EA"


def set_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None):
    run.font.name = "Yu Gothic"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Yu Gothic")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Yu Gothic")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: float = 11, color: RGBColor = INK, bold: bool | None = None):
    for run in paragraph.runs:
        set_font(run, size=size, color=color, bold=bold)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "DADCE0", size: str = "4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def setup_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title_block(doc: Document):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Selenia Aroma - Current Application Specification"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_font(header, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Selenia Aroma | Confidential working specification"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(footer, size=9, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Selenia Aroma")
    set_font(r, size=11, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("現行アプリ仕様書")
    set_font(r, size=26, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("購入者向けWeb/PWA、管理者画面、事業者向けカルテ画面の現時点仕様")
    set_font(r, size=12.5, color=MUTED)

    rows = [
        ("作成日", "2026年7月30日"),
        ("対象", "Selenia Aroma 現行Webアプリ一式"),
        ("購入者向け公開URL", "https://selenia-aroma-customer.vercel.app"),
        ("管理者・事業者側URL", "https://aroma-records-pwa.vercel.app/operator"),
        ("ローカル確認URL", "http://localhost:3100/dashboard"),
        ("現在の公開状態", "購入者向けURLは一時的にログイン入力なしで閲覧可能"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [1.7, 4.8])
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        left.text = label
        right.text = value
        set_cell_shading(left, TABLE_FILL)
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_margins(cell)
            for para in cell.paragraphs:
                set_paragraph_font(para, size=10.5, color=INK, bold=(cell is left))

    add_callout(
        doc,
        "重要な切り分け",
        "購入者向けアプリと管理者・事業者向けアプリは別URL・別用途として扱う。購入者向けUIは、分量・滴数・μLなどの詳細数値を表示しない。管理者・事業者側では制作管理のため、詳細分量を扱う。",
    )


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_border(cell, CALLOUT_BORDER, "8")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, size=10.5, color=INK)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), color=BLUE if level < 3 else DARK_BLUE, bold=True)


def add_para(doc: Document, text: str):
    p = doc.add_paragraph(text)
    set_paragraph_font(p, size=11, color=INK)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, size=10.8, color=INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, TABLE_FILL)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            set_paragraph_font(p, size=10, color=INK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[idx].paragraphs:
                set_paragraph_font(p, size=9.6, color=INK)
    doc.add_paragraph()
    return table


def add_screenshot_section(doc: Document):
    doc.add_page_break()
    add_heading(doc, "14. 画面スクリーンショット", 1)
    add_para(doc, "以下は購入者向け公開URLから、スマートフォン表示を想定して撮影した現行デモ画面である。現在の /login は自動入場仕様のため、個別のログイン画面ではなくホームへ遷移する。")

    for index in range(0, len(SCREENSHOTS), 2):
        if index > 0:
            doc.add_page_break()
        pair = SCREENSHOTS[index:index + 2]
        table = doc.add_table(rows=1, cols=2)
        set_table_width(table, [3.13, 3.13])
        for cell_idx, shot in enumerate(pair):
            title, route, filename, note = shot
            cell = table.rows[0].cells[cell_idx]
            set_cell_border(cell, "ECE6F3", "6")
            set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

            title_p = cell.paragraphs[0]
            title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_p.paragraph_format.space_after = Pt(2)
            run = title_p.add_run(title)
            set_font(run, size=10.5, color=DARK_BLUE, bold=True)

            route_p = cell.add_paragraph()
            route_p.paragraph_format.space_after = Pt(4)
            run = route_p.add_run(route)
            set_font(run, size=8.5, color=MUTED)

            image_p = cell.add_paragraph()
            image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_p.paragraph_format.space_after = Pt(5)
            image_path = SCREENSHOT_DIR / filename
            if image_path.exists():
                image_p.add_run().add_picture(str(image_path), width=Inches(2.6))
            else:
                run = image_p.add_run(f"スクリーンショット未生成: {filename}")
                set_font(run, size=9, color=MUTED)

            note_p = cell.add_paragraph()
            note_p.paragraph_format.space_after = Pt(0)
            run = note_p.add_run(note)
            set_font(run, size=8.7, color=INK)
        doc.add_paragraph()


def add_spec_content(doc: Document):
    add_heading(doc, "1. 概要", 1)
    add_para(doc, "本仕様書は、2026年7月30日時点のSelenia Aroma関連Webアプリの現行仕様を整理したものである。先方共有、要件確認、本番化前の開発整理に利用する。")
    add_para(doc, "現在はNext.jsベースのWeb/PWAとして実装されており、購入者向け公開アプリ、管理者画面、事業者向けカルテ画面を同一コードベースから用途別に切り分けている。")

    add_heading(doc, "2. アプリの切り分け", 1)
    add_table(
        doc,
        ["区分", "用途", "現行URL", "現在の扱い"],
        [
            ["購入者向けアプリ", "購入者が香りの記録、ベースブレンド、精油情報、再購入導線を見る", "https://selenia-aroma-customer.vercel.app", "一時的にログイン不要。/admin と /operator は /dashboard に戻す"],
            ["管理者画面", "顧客とアロマ記録を管理する", "/admin", "本番では管理者ログイン前提。購入者向け公開URLでは非公開扱い"],
            ["事業者向けカルテ", "脳波画像、ヒアリング、配合、制作履歴を業務管理する", "https://aroma-records-pwa.vercel.app/operator", "購入者向けとは別アプリ・別URLとして扱う"],
        ],
        [1.25, 2.0, 1.9, 1.35],
    )
    add_callout(doc, "現在の重要ルール", "購入者向けアプリのUIは現状維持。管理者・事業者側は別アプリとして扱い、購入者向けの表示変更を管理者側に反映しない。")

    add_heading(doc, "3. 現在の公開・ログイン仕様", 1)
    add_bullets(
        doc,
        [
            "購入者向け公開URLでは、現時点のデモ確認を優先し、ログイン情報を入力しなくても閲覧できる。",
            "購入者向けURLで /login を開いた場合、自動的に /dashboard へ遷移する。",
            "内部的にはゲスト顧客セッションを作成し、デモ顧客として表示する。",
            "デモ顧客の表示名は「花咲 美月さん」。",
            "将来的な本番運用ではSupabase Authのメール/パスワードログインへ戻す。",
        ],
    )

    add_heading(doc, "4. 購入者向けアプリ仕様", 1)
    add_para(doc, "購入者向けアプリはスマートフォンファーストのWeb/PWAであり、購入者が自分の香りの記録を気軽に眺め、気分・目的から香りを探し、必要に応じて再購入導線へ進むためのアプリである。")
    add_heading(doc, "4.1 実装済み画面", 2)
    add_table(
        doc,
        ["画面", "URL", "主な役割"],
        [
            ["ログイン/自動入場", "/login", "現デモでは入力なしで /dashboard へ遷移"],
            ["ホーム", "/dashboard", "あいさつ、おすすめカード、次アクション、最近の記録"],
            ["アロマ一覧", "/aromas", "自分のアロマ記録をカードで一覧表示"],
            ["アロマ詳細", "/aromas/[id]", "テーマ、ベースブレンド、追加オイル名、メモ、注意事項、再購入導線"],
            ["再購入確認", "/aromas/[id]/reorder", "外部ECへ移動する前の確認"],
            ["気分から探す", "/moods", "目的別に香りを探す入口"],
            ["気分別詳細", "/moods/[mood]", "目的に合う精油と関連記録を表示"],
            ["精油ページ", "/oils/[slug]", "一般精油36種類の説明、相性、注意メモ"],
            ["ベースブレンド一覧", "/base-blends", "12種類のベースブレンドを図鑑形式で表示"],
            ["ベースブレンド詳細", "/base-blends/[id]", "含まれる精油、目的、関連記録を表示"],
            ["プロフィール", "/profile", "表示名、好み、お気に入り、メニュー"],
        ],
        [1.35, 1.65, 3.5],
    )

    add_heading(doc, "4.2 購入者向け表示ルール", 2)
    add_table(
        doc,
        ["項目", "購入者向け表示", "理由"],
        [
            ["ベースブレンドの配合比率", "表示しない", "内部レシピとして保護する"],
            ["ベースブレンド使用量", "表示しない", "購入者には制作数値を見せず、香りの内容確認に集中させる"],
            ["追加オイルの滴数・ml・μL", "表示しない", "数字が難しく見えるため、オイル名のみ表示する"],
            ["商品容量", "再購入確認画面では表示しない", "購入前確認では商品名・価格・外部移動導線を優先する"],
            ["含まれる精油名", "表示する", "香りの構成を安心して確認できるようにする"],
            ["脳波データ連携ID", "現デモでは表示あり", "将来のカルテアプリとの紐づけ確認用"],
        ],
        [1.8, 1.7, 3.0],
    )
    add_callout(doc, "購入者向けの最新変更", "2026年7月30日時点では、購入者側のアロマ詳細画面から数量・滴数・ml・μL表記を外している。管理者・事業者側の制作管理画面では詳細分量を扱う。")

    add_heading(doc, "4.3 購入者向けデータ", 2)
    add_bullets(
        doc,
        [
            "デモ顧客: 花咲 美月さん、さとしさん、あやかさん。",
            "デモアロマ記録: Relax Night Blend、Morning Refresh、Focus Clear Blend、Happy Citrus、Breath Deep Blend、Floral Harmony。",
            "ベースブレンド: 12種類。公開するのはブレンド番号、仮名、含まれる精油、目的カテゴリ。",
            "一般精油図鑑: 36種類。学名、科名、香りのノート、用途、相性、注意メモを保持。",
            "画像未設定時は、やさしいグラデーションのプレースホルダーを表示する。",
        ],
    )

    add_heading(doc, "5. 管理者画面仕様", 1)
    add_para(doc, "管理者画面は、顧客とアロマ制作記録を登録・管理するための画面である。購入者向け公開URLではアクセスさせない。")
    add_table(
        doc,
        ["画面", "URL", "主な機能"],
        [
            ["管理者ダッシュボード", "/admin", "KPI、最近の記録、新規作成導線、クイックメニュー"],
            ["顧客一覧", "/admin/customers", "顧客プロフィールの確認"],
            ["アロマ記録一覧", "/admin/aromas", "管理者向け記録一覧"],
            ["新規アロマ作成", "/admin/aromas/new", "顧客選択、タイトル、ベースブレンド、画像、追加オイル、再購入URL、公開状態"],
            ["アロマ編集", "/admin/aromas/[id]/edit", "新規作成フォームを再利用した編集画面"],
        ],
        [1.5, 1.8, 3.2],
    )
    add_para(doc, "本番化時にはSupabase Authのadminロールで保護し、customerからは閲覧不可にする。")

    add_heading(doc, "6. 事業者向けカルテアプリ仕様", 1)
    add_para(doc, "事業者向けカルテアプリは、購入者向けアプリとは別アプリとして扱う。脳波画像、ヒアリング回答、注意フラグ、配合量、制作履歴をオペレーターが確認するための業務画面である。")
    add_bullets(
        doc,
        [
            "対象URL: https://aroma-records-pwa.vercel.app/operator",
            "主画面: 顧客カルテ、ベースブレンド、精油図鑑。",
            "脳波画像: サンプル波形とアップロードUIを持つ。現デモではブラウザ内state保存。",
            "香り制作記録: 完成量、配合材料、μL/mL表示切替、比率換算を扱う。",
            "ヒアリング: Googleフォーム等の回答連携を想定し、禁忌・注意フラグを表示する。",
            "ベースブレンド内部情報: 通常表示では伏せ、管理者解除時のみ内部比率や運用メモを表示する。",
        ],
    )
    add_callout(doc, "切り分け方針", "購入者側の「分量を見せない」仕様は、事業者側の制作管理には適用しない。事業者側では、正確な制作・再現のためにμL/mL等の詳細数値を保持する。")

    add_heading(doc, "7. データベース設計", 1)
    add_para(doc, "Supabase Database接続を前提とし、現在は以下のテーブル定義とRLSポリシーを用意している。現デモでは主にモックデータを利用している。")
    add_table(
        doc,
        ["テーブル", "用途", "主な項目"],
        [
            ["profiles", "顧客・管理者プロフィール", "user_id, name, avatar_url, role, created_at"],
            ["aroma_records", "アロマ制作記録", "user_id, title, mood, purpose, notes, image, reorder_url, status, made_at"],
            ["aroma_ingredients", "追加オイル情報", "aroma_record_id, name, amount, unit, sort_order"],
            ["favorites", "お気に入り", "user_id, aroma_record_id, created_at"],
            ["mood_categories", "気分カテゴリ", "name, slug, description, image_url, color"],
            ["base_blends", "公開用ベースブレンド", "code, name, public_ingredients, benefits, color"],
            ["base_blend_private_recipes", "内部レシピ", "base_blend_id, internal_ratio, private_note"],
            ["essential_oils", "一般精油図鑑", "slug, name, botanical_name, common_uses, blends_well_with, safety_note"],
        ],
        [1.85, 1.8, 2.85],
    )

    add_heading(doc, "8. セキュリティ・権限", 1)
    add_table(
        doc,
        ["ユーザー種別", "権限"],
        [
            ["customer", "自分のprofile、自分に紐づくpublishedのaroma_records、自分のお気に入りのみ操作可能。内部レシピは閲覧不可。"],
            ["admin", "全顧客、全記録、ベースブレンド、内部レシピ、精油図鑑を管理可能。"],
            ["未ログイン", "本番では /login 以外へアクセス不可。現デモの購入者向けURLのみ一時的に自動入場。"],
        ],
        [1.55, 4.95],
    )
    add_para(doc, "Supabase RLSでは、customerの自分用閲覧制御とadminの全権限管理を分離する。サービスロールキーはクライアントに露出しない。")

    add_heading(doc, "9. 技術構成", 1)
    add_table(
        doc,
        ["項目", "内容"],
        [
            ["フレームワーク", "Next.js 16.2.6 / App Router"],
            ["言語", "TypeScript"],
            ["UI", "Tailwind CSS 4、スマートフォンファースト、下部タブナビ"],
            ["アイコン", "lucide-react"],
            ["認証・DB・Storage", "Supabase Auth / Database / Storage 接続を想定"],
            ["PWA", "public/manifest.json を用意"],
            ["構成分離", "components, hooks, services, types, data, lib に分離"],
            ["起動コマンド", "npm run dev:3100 / npm run build / npm run start:3100"],
        ],
        [1.7, 4.8],
    )

    add_heading(doc, "10. 環境変数・デプロイ", 1)
    add_table(
        doc,
        ["環境変数", "用途", "現行デモ"],
        [
            ["NEXT_PUBLIC_SUPABASE_URL", "Supabase Project URL", "未設定でもデモモードで動作"],
            ["NEXT_PUBLIC_SUPABASE_ANON_KEY", "Supabase anon public key", "未設定でもデモモードで動作"],
            ["NEXT_PUBLIC_ENABLE_DEMO_MODE", "モックログイン/ゲスト閲覧を有効化", "customer公開URLでは true"],
            ["NEXT_PUBLIC_APP_TARGET", "公開対象の切り替え。customer時は /admin と /operator をブロック", "customer"],
        ],
        [2.2, 2.7, 1.6],
    )
    add_bullets(
        doc,
        [
            "購入者向け公開プロジェクト: selenia-aroma-customer。",
            "購入者向けURL: https://selenia-aroma-customer.vercel.app。",
            "管理者・事業者側URL: https://aroma-records-pwa.vercel.app/operator。",
            "ローカル開発ではポート衝突を避けるため、http://localhost:3100 を使用する。",
        ],
    )

    add_heading(doc, "11. 現在の制限", 1)
    add_bullets(
        doc,
        [
            "購入者向け公開URLはデモ確認用で、現在はログインなしで閲覧できる。",
            "本番Supabase環境への完全接続は未完了。",
            "画像アップロードはUI・サービス層の準備段階で、実運用保存はStorage接続後に完成させる。",
            "外部EC URLはデモURLを含む。",
            "管理者の編集・削除、顧客管理、精油図鑑管理は本番化前に仕上げが必要。",
            "事業者カルテアプリの脳波画像・ドラフト保存は現行デモではブラウザ内状態であり、永続化は未実装。",
        ],
    )

    add_heading(doc, "12. 本番化に向けた次の作業", 1)
    add_table(
        doc,
        ["優先度", "作業", "目的"],
        [
            ["高", "Supabase本番プロジェクト作成、schema.sql適用、RLS確認", "認証・DB・権限を本番化する"],
            ["高", "顧客・管理者ユーザー作成、role設定", "customer/adminの正しい振り分け"],
            ["高", "画像Storage保存と商品画像URLの反映", "商品画像と脳波画像を永続化"],
            ["高", "購入者向けログイン復帰", "デモ公開から本番認証へ切り替える"],
            ["中", "管理者編集・削除、精油図鑑/ベースブレンド管理のDB接続", "運用管理を実データ化"],
            ["中", "事業者カルテのGoogleフォーム/API連携", "ヒアリング回答と制作履歴を紐づける"],
            ["低", "Expo/React Native移植準備", "将来のiOS/Androidアプリ化に備える"],
        ],
        [0.8, 2.7, 3.0],
    )

    add_heading(doc, "13. 受け入れ条件", 1)
    add_bullets(
        doc,
        [
            "購入者向けURLで、ログイン入力なしにホーム画面を確認できる。",
            "購入者向け画面にベースブレンド配合比率、使用量、滴数、μLなどの分量が表示されない。",
            "購入者がアロマ記録、ベースブレンド、精油図鑑、再購入導線を確認できる。",
            "購入者向けURLで /admin と /operator に直接アクセスしても /dashboard に戻る。",
            "管理者・事業者側アプリは購入者向けUI変更の影響を受けない。",
            "npm run build が成功する。",
        ],
    )
    add_screenshot_section(doc)


def build_docx():
    OUT_DOCX.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    setup_styles(doc)
    add_title_block(doc)
    add_spec_content(doc)
    doc.save(DOCX_PATH)
    return DOCX_PATH


def render_docx_to_pdf(docx_path: Path):
    OUT_RENDER.mkdir(parents=True, exist_ok=True)
    renderer = Path(r"C:\Users\halcy\.codex\plugins\cache\openai-primary-runtime\documents\26.727.11326\skills\documents\render_docx.py")
    python = Path(r"C:\Users\halcy\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe")
    cmd = [str(python), str(renderer), str(docx_path), "--output_dir", str(OUT_RENDER), "--emit_pdf"]
    subprocess.run(cmd, check=True, capture_output=True, text=True)
    emitted = OUT_RENDER / f"{docx_path.stem}.pdf"
    OUT_PDF.mkdir(parents=True, exist_ok=True)
    shutil.copy2(emitted, PDF_PATH)


class SectionRule(Flowable):
    def __init__(self, color=colors.HexColor("#D8C8EA"), width=6.5 * inch):
        super().__init__()
        self.color = color
        self.width = width
        self.height = 5

    def draw(self):
        self.canv.setStrokeColor(self.color)
        self.canv.setLineWidth(1)
        self.canv.line(0, 2, self.width, 2)


def pdf_styles():
    font_path = Path(r"C:\Windows\Fonts\NotoSansJP-VF.ttf")
    if font_path.exists():
        pdfmetrics.registerFont(TTFont("NotoSansJP", str(font_path)))
        base = "NotoSansJP"
    else:
        pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
        base = "HeiseiKakuGo-W5"
    return {
        "title": ParagraphStyle(
            "title",
            fontName=base,
            fontSize=24,
            leading=32,
            textColor=colors.black,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            fontName=base,
            fontSize=11.5,
            leading=17,
            textColor=colors.HexColor("#555555"),
            spaceAfter=14,
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "h1",
            fontName=base,
            fontSize=15,
            leading=20,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "h2",
            fontName=base,
            fontSize=12,
            leading=17,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=9,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "body",
            fontName=base,
            fontSize=9.5,
            leading=14.5,
            textColor=colors.HexColor("#232323"),
            spaceAfter=5,
            wordWrap="CJK",
        ),
        "bullet": ParagraphStyle(
            "bullet",
            fontName=base,
            fontSize=9.3,
            leading=14.2,
            leftIndent=13,
            firstLineIndent=-8,
            spaceAfter=3,
            wordWrap="CJK",
        ),
        "small": ParagraphStyle(
            "small",
            fontName=base,
            fontSize=8.2,
            leading=11.5,
            textColor=colors.HexColor("#555555"),
            wordWrap="CJK",
        ),
        "callout_title": ParagraphStyle(
            "callout_title",
            fontName=base,
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=3,
            wordWrap="CJK",
        ),
        "callout_body": ParagraphStyle(
            "callout_body",
            fontName=base,
            fontSize=9,
            leading=13.2,
            textColor=colors.HexColor("#232323"),
            wordWrap="CJK",
        ),
        "screenshot_title": ParagraphStyle(
            "screenshot_title",
            fontName=base,
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=1,
            wordWrap="CJK",
        ),
        "screenshot_route": ParagraphStyle(
            "screenshot_route",
            fontName=base,
            fontSize=7.4,
            leading=10,
            textColor=colors.HexColor("#666666"),
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "screenshot_note": ParagraphStyle(
            "screenshot_note",
            fontName=base,
            fontSize=7.2,
            leading=10.5,
            textColor=colors.HexColor("#333333"),
            wordWrap="CJK",
        ),
    }


def add_pdf_table(story, styles, headers, rows, widths):
    data = [[Paragraph(h, styles["small"]) for h in headers]]
    for row in rows:
        data.append([Paragraph(str(value), styles["small"]) for value in row])
    table = Table(data, colWidths=[w * inch for w in widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#DADCE0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(table)
    story.append(Spacer(1, 8))


def add_pdf_callout(story, styles, title, body):
    table = Table(
        [[Paragraph(title, styles["callout_title"]), Paragraph(body, styles["callout_body"])]],
        colWidths=[1.45 * inch, 5.05 * inch],
    )
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F6F2FB")),
        ("BOX", (0, 0), (-1, -1), 0.7, colors.HexColor("#D8C8EA")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(table)
    story.append(Spacer(1, 8))


def add_pdf_bullets(story, styles, items):
    for item in items:
        story.append(Paragraph(f"• {item}", styles["bullet"]))


def add_pdf_screenshot_section(story, styles):
    story.append(PageBreak())
    story.append(Paragraph("13. 画面スクリーンショット", styles["h1"]))
    story.append(Paragraph("以下は購入者向け公開URLから、スマートフォン表示を想定して撮影した現行デモ画面である。現在の /login は自動入場仕様のため、個別のログイン画面ではなくホームへ遷移する。", styles["body"]))

    for index in range(0, len(SCREENSHOTS), 2):
        if index > 0:
            story.append(PageBreak())
        pair = SCREENSHOTS[index:index + 2]
        row = []
        for title, route, filename, note in pair:
            image_path = SCREENSHOT_DIR / filename
            cell_flowables = [
                Paragraph(title, styles["screenshot_title"]),
                Paragraph(route, styles["screenshot_route"]),
            ]
            if image_path.exists():
                cell_flowables.append(RLImage(str(image_path), width=2.48 * inch, height=5.36 * inch))
            else:
                cell_flowables.append(Paragraph(f"スクリーンショット未生成: {filename}", styles["small"]))
            cell_flowables.append(Paragraph(note, styles["screenshot_note"]))
            row.append(cell_flowables)
        while len(row) < 2:
            row.append([])
        table = Table([row], colWidths=[3.16 * inch, 3.16 * inch])
        table.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#ECE6F3")),
            ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#ECE6F3")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]))
        story.append(table)


def build_pdf():
    OUT_PDF.mkdir(parents=True, exist_ok=True)
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.85 * inch,
        bottomMargin=0.85 * inch,
        title="Selenia Aroma 現行アプリ仕様書",
    )
    story = []

    story.append(Paragraph("Selenia Aroma 現行アプリ仕様書", styles["title"]))
    story.append(Paragraph("購入者向けWeb/PWA、管理者画面、事業者向けカルテ画面の現時点仕様 | 作成日: 2026年7月30日", styles["subtitle"]))
    add_pdf_table(
        story,
        styles,
        ["項目", "内容"],
        [
            ["購入者向け公開URL", "https://selenia-aroma-customer.vercel.app"],
            ["管理者・事業者側URL", "https://aroma-records-pwa.vercel.app/operator"],
            ["ローカル確認URL", "http://localhost:3100/dashboard"],
            ["現在の公開状態", "購入者向けURLは一時的にログイン入力なしで閲覧可能"],
            ["デモ顧客名", "花咲 美月さん"],
        ],
        [1.75, 4.75],
    )
    add_pdf_callout(
        story,
        styles,
        "重要な切り分け",
        "購入者向けアプリと管理者・事業者向けアプリは別URL・別用途として扱う。購入者側には分量・滴数・μLなどの詳細数値を表示しない。管理者・事業者側では制作管理のため詳細分量を扱う。",
    )

    story.append(Paragraph("1. 概要", styles["h1"]))
    story.append(Paragraph("本仕様書は、2026年7月30日時点のSelenia Aroma関連Webアプリの現行仕様を整理したものである。先方共有、要件確認、本番化前の開発整理に利用する。", styles["body"]))
    story.append(Paragraph("現在はNext.jsベースのWeb/PWAとして実装されており、購入者向け公開アプリ、管理者画面、事業者向けカルテ画面を同一コードベースから用途別に切り分けている。", styles["body"]))

    story.append(Paragraph("2. アプリの切り分け", styles["h1"]))
    add_pdf_table(
        story,
        styles,
        ["区分", "用途", "現行URL/扱い"],
        [
            ["購入者向けアプリ", "購入者が香りの記録、ベースブレンド、精油情報、再購入導線を見る", "https://selenia-aroma-customer.vercel.app。一時的にログイン不要。/admin と /operator は /dashboard に戻す。"],
            ["管理者画面", "顧客とアロマ記録を管理する", "/admin。本番では管理者ログイン前提。購入者向け公開URLでは非公開扱い。"],
            ["事業者向けカルテ", "脳波画像、ヒアリング、配合、制作履歴を業務管理する", "https://aroma-records-pwa.vercel.app/operator。購入者向けとは別アプリ・別URL。"],
        ],
        [1.35, 2.1, 3.05],
    )
    add_pdf_callout(story, styles, "現在の重要ルール", "購入者向けアプリのUIは現状維持。管理者・事業者側は別アプリとして扱い、購入者向けの表示変更を管理者側に反映しない。")

    story.append(Paragraph("3. 現在の公開・ログイン仕様", styles["h1"]))
    add_pdf_bullets(story, styles, [
        "購入者向け公開URLでは、現時点のデモ確認を優先し、ログイン情報を入力しなくても閲覧できる。",
        "購入者向けURLで /login を開いた場合、自動的に /dashboard へ遷移する。",
        "内部的にはゲスト顧客セッションを作成し、デモ顧客として表示する。",
        "将来的な本番運用ではSupabase Authのメール/パスワードログインへ戻す。",
    ])

    story.append(Paragraph("4. 購入者向けアプリ仕様", styles["h1"]))
    story.append(Paragraph("購入者向けアプリはスマートフォンファーストのWeb/PWAであり、購入者が自分の香りの記録を気軽に眺め、気分・目的から香りを探し、必要に応じて再購入導線へ進むためのアプリである。", styles["body"]))
    add_pdf_table(
        story,
        styles,
        ["画面", "URL", "主な役割"],
        [
            ["ログイン/自動入場", "/login", "現デモでは入力なしで /dashboard へ遷移"],
            ["ホーム", "/dashboard", "あいさつ、おすすめカード、次アクション、最近の記録"],
            ["アロマ一覧", "/aromas", "自分のアロマ記録をカードで一覧表示"],
            ["アロマ詳細", "/aromas/[id]", "テーマ、ベースブレンド、追加オイル名、メモ、注意事項、再購入導線"],
            ["再購入確認", "/aromas/[id]/reorder", "外部ECへ移動する前の確認"],
            ["気分から探す", "/moods", "目的別に香りを探す入口"],
            ["精油ページ", "/oils/[slug]", "一般精油36種類の説明、相性、注意メモ"],
            ["ベースブレンド", "/base-blends", "12種類のベースブレンドを図鑑形式で表示"],
            ["プロフィール", "/profile", "表示名、好み、お気に入り、メニュー"],
        ],
        [1.35, 1.7, 3.45],
    )

    story.append(Paragraph("5. 購入者向け表示ルール", styles["h1"]))
    add_pdf_table(
        story,
        styles,
        ["項目", "購入者向け表示", "理由"],
        [
            ["ベースブレンド配合比率", "表示しない", "内部レシピとして保護"],
            ["ベースブレンド使用量", "表示しない", "制作数値を見せず香りの内容確認に集中させる"],
            ["追加オイルの滴数・ml・μL", "表示しない", "数字が難しく見えるため、オイル名のみ表示"],
            ["商品容量", "再購入確認画面では表示しない", "商品名・価格・外部移動導線を優先"],
            ["含まれる精油名", "表示する", "香りの構成を安心して確認できるようにする"],
        ],
        [1.75, 1.8, 2.95],
    )

    story.append(Paragraph("6. 管理者画面仕様", styles["h1"]))
    add_pdf_bullets(story, styles, [
        "管理者画面は、顧客とアロマ制作記録を登録・管理するための画面である。",
        "主な画面は /admin、/admin/customers、/admin/aromas、/admin/aromas/new、/admin/aromas/[id]/edit。",
        "入力項目は顧客選択、タイトル、テーマ、ベースブレンド、商品画像、追加オイル、再購入URL、公開状態など。",
        "本番化時にはSupabase Authのadminロールで保護し、customerからは閲覧不可にする。",
    ])

    story.append(Paragraph("7. 事業者向けカルテアプリ仕様", styles["h1"]))
    add_pdf_bullets(story, styles, [
        "対象URLは https://aroma-records-pwa.vercel.app/operator。",
        "主画面は顧客カルテ、ベースブレンド、精油図鑑。",
        "脳波画像、Googleフォーム想定のヒアリング回答、禁忌・注意フラグ、制作履歴を一元管理する。",
        "香り制作記録では完成量、配合材料、μL/mL表示切替、比率換算を扱う。",
        "購入者側の「分量を見せない」仕様は、事業者側の制作管理には適用しない。",
    ])

    story.append(Paragraph("8. データベース設計", styles["h1"]))
    add_pdf_table(
        story,
        styles,
        ["テーブル", "用途", "主な項目"],
        [
            ["profiles", "顧客・管理者プロフィール", "user_id, name, avatar_url, role, created_at"],
            ["aroma_records", "アロマ制作記録", "title, mood, purpose, notes, image, reorder_url, status, made_at"],
            ["aroma_ingredients", "追加オイル情報", "name, amount, unit, sort_order"],
            ["favorites", "お気に入り", "user_id, aroma_record_id, created_at"],
            ["base_blends", "公開用ベースブレンド", "code, name, public_ingredients, benefits, color"],
            ["base_blend_private_recipes", "内部レシピ", "internal_ratio, private_note"],
            ["essential_oils", "一般精油図鑑", "slug, name, botanical_name, blends_well_with, safety_note"],
        ],
        [1.75, 1.65, 3.1],
    )

    story.append(Paragraph("9. セキュリティ・権限", styles["h1"]))
    add_pdf_bullets(story, styles, [
        "customerは自分のprofile、自分に紐づくpublishedのaroma_records、自分のお気に入りのみ操作可能。",
        "adminは全顧客、全記録、ベースブレンド、内部レシピ、精油図鑑を管理可能。",
        "本番では未ログインユーザーは/login以外へアクセス不可。現デモの購入者向けURLのみ一時的に自動入場。",
        "Supabase RLSでcustomer/adminの権限を分離し、サービスロールキーはクライアントに露出しない。",
    ])

    story.append(Paragraph("10. 技術構成・環境変数", styles["h1"]))
    add_pdf_table(
        story,
        styles,
        ["項目", "内容"],
        [
            ["フレームワーク", "Next.js 16.2.6 / App Router"],
            ["言語/UI", "TypeScript / Tailwind CSS 4 / lucide-react"],
            ["認証・DB・Storage", "Supabase Auth / Database / Storage 接続を想定"],
            ["PWA", "public/manifest.json を用意"],
            ["環境変数", "NEXT_PUBLIC_SUPABASE_URL, NEXT_PUBLIC_SUPABASE_ANON_KEY, NEXT_PUBLIC_ENABLE_DEMO_MODE, NEXT_PUBLIC_APP_TARGET"],
            ["起動", "npm run dev:3100 / npm run build / npm run start:3100"],
        ],
        [1.85, 4.65],
    )

    story.append(Paragraph("11. 現在の制限と本番化タスク", styles["h1"]))
    add_pdf_bullets(story, styles, [
        "購入者向け公開URLはデモ確認用で、現在はログインなしで閲覧できる。",
        "本番Supabase環境への完全接続、Storage保存、管理者編集・削除、外部EC URL整備が必要。",
        "事業者カルテアプリの脳波画像・ドラフト保存は現行デモではブラウザ内状態であり、永続化は未実装。",
        "本番化時はSupabase本番プロジェクト作成、schema.sql適用、RLS確認、ユーザー作成、購入者向けログイン復帰を行う。",
    ])

    story.append(Paragraph("12. 受け入れ条件", styles["h1"]))
    add_pdf_bullets(story, styles, [
        "購入者向けURLで、ログイン入力なしにホーム画面を確認できる。",
        "購入者向け画面にベースブレンド配合比率、使用量、滴数、μLなどの分量が表示されない。",
        "購入者がアロマ記録、ベースブレンド、精油図鑑、再購入導線を確認できる。",
        "購入者向けURLで /admin と /operator に直接アクセスしても /dashboard に戻る。",
        "管理者・事業者側アプリは購入者向けUI変更の影響を受けない。",
        "npm run build が成功する。",
    ])
    add_pdf_screenshot_section(story, styles)

    def footer(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont(styles["small"].fontName, 7.5)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawString(inch, 0.45 * inch, "Selenia Aroma | Current Application Specification")
        canvas.drawRightString(letter[0] - inch, 0.45 * inch, f"Page {doc_obj.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def rasterize_pdf():
    OUT_RENDER.mkdir(parents=True, exist_ok=True)
    for old_page in OUT_RENDER.glob("pdf-page-*.png"):
        old_page.unlink()
    prefix = OUT_RENDER / "pdf-page"
    pdftoppm = Path(r"C:\Users\halcy\.cache\codex-runtimes\codex-primary-runtime\dependencies\native\poppler\Library\bin\pdftoppm.exe")
    subprocess.run([str(pdftoppm), "-png", str(PDF_PATH), str(prefix)], check=True)


if __name__ == "__main__":
    path = build_docx()
    try:
        render_docx_to_pdf(path)
    except Exception as exc:
        print(f"DOCX render unavailable, generating PDF directly: {exc}")
        build_pdf()
    rasterize_pdf()
    print(path)
    print(PDF_PATH)
